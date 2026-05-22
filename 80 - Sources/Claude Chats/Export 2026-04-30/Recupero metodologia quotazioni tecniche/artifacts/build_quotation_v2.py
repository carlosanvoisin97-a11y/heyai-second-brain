#!/usr/bin/env python3
"""
Recreate NOT quotation with Eye Cookies template formatting.
Uses the EC template as base to preserve headers, footers, logos, styles.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

# ── Open template ──
doc = Document('/mnt/user-data/uploads/NOT_Quotazione_Eye_Cookies_def.docx')

# ── Clear all body content ──
body = doc.element.body
# Remove all children except sectPr (section properties - keeps header/footer refs)
children_to_remove = []
for child in body:
    if child.tag != qn('w:sectPr'):
        children_to_remove.append(child)
for child in children_to_remove:
    body.remove(child)

# ── Formatting helpers ──
NAVY = RGBColor(0x1F, 0x4E, 0x79)
BLUE = RGBColor(0x2E, 0x86, 0xAB)
BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN_TEXT = RGBColor(0x27, 0xAE, 0x60)

def section_title(text):
    """Main section title: Bold 14pt Aptos Navy - matches EC style"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Aptos'
    run.font.color.rgb = NAVY
    p.space_after = Pt(6)
    p.space_before = Pt(18)
    return p

def subsection_title(text):
    """Subsection title: Bold 12pt Aptos Blue accent - matches EC H2"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Aptos'
    run.font.color.rgb = BLUE
    p.space_after = Pt(4)
    p.space_before = Pt(14)
    return p

def sub3_title(text):
    """Sub-subsection: Bold 11pt Aptos Black"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Aptos'
    run.font.color.rgb = BLACK
    p.space_after = Pt(4)
    p.space_before = Pt(10)
    return p

def body_text(text):
    """Normal body paragraph"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Aptos'
    run.font.size = Pt(11)
    p.space_after = Pt(6)
    return p

def bold_value(label, value):
    """Bold label + normal value on same paragraph"""
    p = doc.add_paragraph()
    r1 = p.add_run(label)
    r1.bold = True
    r1.font.name = 'Aptos'
    r1.font.size = Pt(11)
    r2 = p.add_run(value)
    r2.font.name = 'Aptos'
    r2.font.size = Pt(11)
    p.space_after = Pt(4)
    return p

def note_text(text, prefix="Nota: "):
    """Bold prefix + italic text for notes"""
    p = doc.add_paragraph()
    r1 = p.add_run(prefix)
    r1.bold = True
    r1.font.name = 'Aptos'
    r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.italic = True
    r2.font.name = 'Aptos'
    r2.font.size = Pt(11)
    p.space_after = Pt(6)
    return p

def empty_para():
    p = doc.add_paragraph()
    p.space_after = Pt(0)
    p.space_before = Pt(0)
    return p

def set_cell_shading(cell, color):
    """Set cell background color"""
    tc = cell._element
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}/>') 
        tc.insert(0, tcPr)
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    # Remove existing shading
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    tcPr.append(shd)

def set_cell_width(cell, width_cm):
    tc = cell._element
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}/>') 
        tc.insert(0, tcPr)
    w = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{int(width_cm * 567)}" w:type="dxa"/>')
    for old in tcPr.findall(qn('w:tcW')):
        tcPr.remove(old)
    tcPr.append(w)

def format_cell(cell, text, bold=False, font_size=11, font_color=None, bg=None):
    """Format a table cell in EC style"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.space_after = Pt(2)
    p.space_before = Pt(2)
    run = p.add_run(str(text))
    run.font.name = 'Aptos'
    run.font.size = Pt(font_size)
    if bold:
        run.bold = True
    if font_color:
        run.font.color.rgb = font_color
    if bg:
        set_cell_shading(cell, bg)

def make_ec_table(headers, rows, col_widths_cm=None):
    """
    Create a table matching Eye Cookies style.
    headers: list of strings
    rows: list of lists. Each cell can be str or dict with keys: text, bold, bg, color, font_size
    col_widths_cm: list of column widths in cm
    """
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Style the table with borders
    tbl = table._element
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>') 
        tbl.insert(0, tblPr)
    
    # Set table borders
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '</w:tblBorders>'
    )
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    tblPr.append(borders)
    
    # Header row
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        format_cell(cell, h, bold=True, font_size=11, font_color=WHITE, bg="1F4E79")
        if col_widths_cm:
            set_cell_width(cell, col_widths_cm[ci])
    
    # Data rows with alternating shading
    for ri, row_data in enumerate(rows):
        alt_bg = "F2F2F2" if ri % 2 == 1 else None
        for ci, cell_data in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            if isinstance(cell_data, dict):
                text = cell_data.get('text', '')
                bold = cell_data.get('bold', False)
                bg = cell_data.get('bg', alt_bg)
                color = cell_data.get('color', None)
                fs = cell_data.get('font_size', 11)
                format_cell(cell, text, bold=bold, font_size=fs, font_color=color, bg=bg)
            else:
                format_cell(cell, str(cell_data), bg=alt_bg)
            if col_widths_cm:
                set_cell_width(cell, col_widths_cm[ci])
    
    return table

def total_row(table, row_idx, bg="D9E2F3"):
    """Mark a row as a total row"""
    for cell in table.rows[row_idx].cells:
        set_cell_shading(cell, bg)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

def green_cell(cell, text):
    """Green highlight cell like EC template"""
    format_cell(cell, text, bold=True, font_color=GREEN_TEXT, bg="D5F5E3")


# ════════════════════════════════════════════════════════════════
# ═══  DOCUMENT CONTENT  ═══════════════════════════════════════
# ════════════════════════════════════════════════════════════════

# ── TITLE ──
p = doc.add_paragraph()
p.style = doc.styles['Heading 1']
p.space_before = Pt(0)

p = doc.add_paragraph()
p.style = doc.styles['Heading 1']
run = p.add_run("Proposta di Investimento: Piattaforma NOT")
run.bold = True
run.font.color.rgb = RGBColor(0x15, 0x60, 0x82)
p.space_after = Pt(12)

body_text("Modulo Logistica & Stock Management + Modulo Gestione Vendite — Fase 1 (MVP)")

# ════════════════════════════════════════════════════════════════
# 1. OGGETTO DELLA PROPOSTA
# ════════════════════════════════════════════════════════════════
section_title("1. Oggetto della Proposta")

body_text("La presente proposta riguarda la progettazione, lo sviluppo e l'implementazione dei primi due moduli della piattaforma AI per NOT (NoLoop on Trade): il Modulo Logistica & Stock Management e il Modulo Gestione Vendite.")

body_text("NOT gestisce oggi circa 30-33 pop-up store in centri commerciali italiani con un team Operations di 20 persone che opera attraverso processi manuali frammentati su Excel, Google Sheets e WhatsApp. L'inefficienza operativa limita la scalabilità e genera errori sistematici nella gestione di stock, vendite e reportistica.")

body_text("La piattaforma sostituisce i processi manuali con un sistema integrato che include: gestione predittiva degli ordini basata su AI, tracciamento stock end-to-end (magazzini e pop-up), form unico di registrazione vendite (sostitutivo della tripla compilazione attuale), cross-check automatici, dashboard real-time, timbratura digitale e analytics.")

# ════════════════════════════════════════════════════════════════
# 2. SOLUZIONE TECNICA
# ════════════════════════════════════════════════════════════════
section_title("2. Soluzione Tecnica")

body_text("Lo sviluppo comprende i seguenti moduli funzionali:")

empty_para()

# Module table
make_ec_table(
    ["Modulo", "Descrizione"],
    [
        [{"text": "1. Setup Iniziale (F0)", "bold": True}, "Interfaccia censimento giacenze iniziali per ogni pop-up e magazzino (Milano e Roma). Import storico vendite 5 anni per alimentazione modello predittivo."],
        [{"text": "2. Modello Predittivo AI (F1)", "bold": True}, "AI basata su storico vendite, stagionalità e trend. Genera proposte ordine a BAT (aggregato) e proposte replenishment per singolo pop-up. Gestione ereditarietà profilo vendita per nuovi prodotti. Parametri regolabili per calibrazione. Ottimizzazione distribuzione in caso di stock insufficiente."],
        [{"text": "3. Workflow Approvazione (F2)", "bold": True}, "Visualizzazione, modifica e approvazione ordini (Carmen/Laila). Conversione ordine approvato in task per magazziniere. Vista Area Manager in sola lettura."],
        [{"text": "4. Magazzino & Pop-up (F3-F4)", "bold": True}, "Check-out magazzino con conferma quantità aggregate e alert discrepanza. Check-in merce da BAT con verifica bolla. Check-in pop-up con notifica spedizione, conferma ricezione, registrazione codici myglow e aggiornamento automatico giacenze."],
        [{"text": "5. Form Vendita Unico (F5)", "bold": True}, "Scansione codice myglow tramite fotocamera con riconoscimento AI (modello/colore/prezzo). Auto-compilazione operatore e timestamp. 3 flussi vendita configurabili (scansione obbligatoria, inserimento manuale, vendita Shopify). Gestione prezzi centralizzata con logica sconto clienti BAT. Registrazione in due step. Cross-check triplo automatico (vendite/stock/incassi)."],
        [{"text": "6. Vendita Indiretta & Timbratura (F6+T)", "bold": True}, "Tipologia pop-up configurabile (vendita diretta vs tabaccaio). Tracciamento legame vendita NOT→tabaccaio. Timbratura digitale check-in/check-out con geolocalizzazione. Conteggio ore automatico e correlazione vendite/ora."],
        [{"text": "7. Dashboard & Analytics (D)", "bold": True}, "Stock real-time per pop-up/magazzino/prodotto. Previsione rottura stock. Segnalazione anomalie consumo. Performance per pop-up e venditore. Regolazione parametri modello predittivo. Export Excel/CSV. Report settimanale BAT auto-generato."],
    ],
    [5.5, 12.0]
)

empty_para()

body_text("Il progetto include inoltre: architettura multi-tenant predisposta per supporto futuri clienti (oltre GLO/BAT), gestione pop-up con categorie prodotto diverse, Web App mobile-first (PWA) con login identificativo, 4 profili utente (Admin, Magazziniere, Area Manager, Promoter/Steward), interfaccia magazziniere semplificata e gestione connettività scarsa con interfaccia di re-inserimento.")

# ════════════════════════════════════════════════════════════════
# 3. PIANO DI LAVORO
# ════════════════════════════════════════════════════════════════
section_title("3. Piano di Lavoro")

body_text("Lo sviluppo seguirà un approccio agile articolato in 4 fasi:")

empty_para()

make_ec_table(
    ["Fase", "Durata", "Deliverable"],
    [
        ["Discovery & Design", "4 settimane", "Process mapping, architettura tecnica, wireframe interfacce, backlog approvato, sprint planning"],
        ["Sviluppo Core", "10 settimane", "Setup infrastruttura e autenticazione, form vendita con AI, check-in/check-out, modello predittivo, workflow approvazione, cross-check, timbratura digitale"],
        ["Dashboard, Test & Go-Live", "6 settimane", "Dashboard analytics e report, UAT con team NOT, bug fixing, deployment, training utenti, go-live su pop-up pilota"],
        ["Hypercare & Stabilizzazione", "4 settimane", "Supporto post go-live (8h/sett), rollout progressivo, fine-tuning modello predittivo su dati reali"],
        [{"text": "TOTALE", "bold": True, "bg": "D9E2F3"}, {"text": "24 settimane", "bold": True, "bg": "D9E2F3"}, {"text": "MVP completo e operativo su tutti i pop-up", "bold": True, "bg": "D9E2F3"}],
    ],
    [4.5, 3.0, 10.0]
)

# ════════════════════════════════════════════════════════════════
# 4. INVESTIMENTO
# ════════════════════════════════════════════════════════════════
section_title("4. Investimento")

empty_para()

make_ec_table(
    ["Pacchetto", "Scope", "Importo"],
    [
        ["Base", "Core MVP: logistica + vendite senza AI avanzata e senza integrazioni esterne", {"text": "€ 135.000", "bold": True}],
        [{"text": "Recommended", "bold": True, "bg": "D9E2F3"}, {"text": "MVP completo: AI predittiva, form unico, dashboard, timbratura digitale, integrazione Shopify, hypercare 4 settimane", "bg": "D9E2F3"}, {"text": "€ 168.000", "bold": True, "bg": "D9E2F3"}],
        ["Advanced", "MVP + integrazioni BAT/Pard (API o RPA) + calendario turni + OCR scontrino + chat AI + hypercare estesa", {"text": "€ 215.000", "bold": True}],
    ],
    [3.5, 10.5, 3.5]
)

empty_para()

subsection_title("4.1 Breakdown per Fase (Pacchetto Recommended)")

make_ec_table(
    ["Fase", "Importo"],
    [
        ["Discovery & Design (sett. 1-4)", "€ 28.000"],
        ["Sviluppo Core (sett. 5-14)", "€ 92.000"],
        ["Dashboard, Test & Go-Live (sett. 15-20)", "€ 38.000"],
        ["Hypercare & Stabilizzazione (sett. 21-24)", "€ 10.000"],
        [{"text": "TOTALE PACCHETTO RECOMMENDED", "bold": True, "bg": "D9E2F3"}, {"text": "€ 168.000", "bold": True, "bg": "D9E2F3", "font_size": 13}],
    ],
    [12.0, 5.5]
)

empty_para()

subsection_title("4.2 Breakdown per Modulo (Pacchetto Recommended)")

make_ec_table(
    ["Componente", "Importo", "% sul totale"],
    [
        ["Modulo Logistica & Stock Management", "€ 62.000", "37%"],
        ["Modulo Gestione Vendite", "€ 72.000", "43%"],
        ["Dashboard & Analytics (trasversale)", "€ 18.000", "11%"],
        ["Project Management, Training, Hypercare", "€ 16.000", "9%"],
        [{"text": "TOTALE", "bold": True, "bg": "D9E2F3", "font_size": 12}, {"text": "€ 168.000", "bold": True, "bg": "D9E2F3", "font_size": 13}, {"text": "100%", "bold": True, "bg": "D9E2F3"}],
    ],
    [8.5, 5.0, 4.0]
)

empty_para()

bold_value("Incluso: ", "Codice sorgente completo, documentazione tecnica e guida operativa, training iniziale per Admin e key user (2 sessioni), 4 settimane di hypercare post-lancio (supporto dedicato 8h/settimana), report di collaudo, modello AI predittivo addestrato su storico 5 anni.")

# ════════════════════════════════════════════════════════════════
# 5. MODALITA DI PAGAMENTO
# ════════════════════════════════════════════════════════════════
section_title("5. Modalità di Pagamento")

empty_para()

t = make_ec_table(
    ["Milestone", "Importo"],
    [
        ["30% alla firma del contratto", "€ 50.400"],
        ["20% al completamento Discovery (sett. 4) — backlog approvato e architettura validata", "€ 33.600"],
        ["30% alla demo intermedia (sett. 14) — core funzionante in staging", "€ 50.400"],
        ["20% al go-live e sign-off UAT (sett. 20) — rilascio in produzione", "€ 33.600"],
    ],
    [12.0, 5.5]
)

# ════════════════════════════════════════════════════════════════
# 6. RAZIONALE ECONOMICO
# ════════════════════════════════════════════════════════════════
section_title("6. Razionale Economico")

sub3_title("Premessa")

body_text("Il business case documentato identifica un risparmio operativo significativo derivante dall'eliminazione della riconciliazione tripla dei dati di vendita e dall'automazione degli ordini logistici. Il saving proviene prevalentemente dalla riallocazione di capacità (stesse risorse, attività a maggior valore), con possibile riduzione headcount per i 2 Data Analyst oggi full-time sulla riconciliazione. La piattaforma consentirebbe a NOT di gestire una crescita da 30-33 a 50+ pop-up senza assumere personale aggiuntivo.")

subsection_title("6.1 Coerenza Investimento / Saving")

make_ec_table(
    ["METRICA", "VALORE"],
    [
        [{"text": "Saving annuo a regime (realistico)", "bold": True}, "€ 214.268"],
        [{"text": "Saving cumulato 3 anni (realistico)", "bold": True}, "€ 557.097"],
        [{"text": "Investimento Recommended", "bold": True}, {"text": "€ 168.000", "bold": True}],
        [{"text": "Payback (mesi dalla messa in produzione)", "bold": True}, {"text": "~ 9-10 mesi", "bold": True, "color": GREEN_TEXT, "bg": "D5F5E3"}],
        [{"text": "ROI a 3 anni", "bold": True}, {"text": "3,3x", "bold": True, "color": GREEN_TEXT, "bg": "D5F5E3"}],
        [{"text": "% EBITDA annuo assorbito (su €1,44-1,80M)", "bold": True}, "9-12%"],
        [{"text": "Collocazione nel range affordability (€150-270K Fase 1)", "bold": True}, "Entro range realistico"],
    ],
    [11.0, 6.5]
)

empty_para()

body_text("L'investimento Recommended si posiziona nel cuore del range di Fase 1 dello scenario realistico di affordability (€ 150.000-270.000), con un impatto sull'EBITDA annuo contenuto (9-12%) e un payback inferiore a 12 mesi. Il rapporto investimento/saving cumulato a 3 anni (30%) è ampiamente nella fascia attraente (30-50% del saving) identificata nel business case.")

subsection_title("6.2 Split Saving per Modulo")

make_ec_table(
    ["Modulo", "Saving annuo regime", "% sul totale"],
    [
        ["Gestione Vendite", "€ 124.870", "58%"],
        ["Logistica", "€ 89.398", "42%"],
        [{"text": "TOTALE", "bold": True, "bg": "D9E2F3"}, {"text": "€ 214.268", "bold": True, "bg": "D9E2F3"}, {"text": "100%", "bold": True, "bg": "D9E2F3"}],
    ],
    [8.0, 5.0, 4.5]
)

empty_para()

note_text("I due moduli sono fortemente interconnessi (la logistica dipende dai dati di vendita). Implementare solo un modulo riduce il saving complessivo ma anche la complessità progettuale. Si raccomanda l'implementazione congiunta per massimizzare il valore.")

# ════════════════════════════════════════════════════════════════
# 7. SCOPE DETTAGLIATO E DELIVERABLE
# ════════════════════════════════════════════════════════════════
section_title("7. Scope Dettagliato e Deliverable")

subsection_title("7.1 Modulo Logistica & Stock Management — Funzionalità Incluse MVP")

body_text("F0 — Setup iniziale: Interfaccia censimento giacenze iniziali per ogni pop-up e magazzino (Milano e Roma). Import storico vendite 5 anni per alimentazione modello predittivo.")

body_text("F1 — Modello predittivo e generazione ordini: AI basata su storico vendite, stagionalità e trend. Genera proposte ordine a BAT (aggregato) e proposte replenishment per singolo pop-up. Gestione ereditarietà profilo vendita per nuovi prodotti. Parametri regolabili per calibrazione modello. Ottimizzazione distribuzione in caso di stock insufficiente.")

body_text("F2 — Workflow approvazione: Visualizzazione, modifica e approvazione ordini (Carmen/Laila). Conversione ordine approvato in task per magazziniere. Vista Area Manager in sola lettura.")

body_text("F3 — Check-out magazzino: Conferma quantità caricate per spedizione (aggregate, non per singolo prodotto). Alert discrepanza automatico. Aggiornamento giacenze magazzino. Check-in merce da BAT con verifica bolla.")

body_text("F4 — Check-in pop-up: Notifica spedizione in arrivo con dettaglio prodotti. Conferma ricezione con alert discrepanza. Registrazione codici myglow al check-in (inserimento manuale alfanumerico). Aggiornamento automatico giacenze pop-up.")

subsection_title("7.2 Modulo Logistica — Escluso / Future Option")

body_text("Integrazione Google Trends per pattern stagionali (evoluzione futura). Scansione barcode prodotto per prodotto a magazzino (esclusa per decisione operativa). Integrazione diretta con sistemi BAT per visibilità stock fornitore (da valutare come evolutiva). Gestione rapporti commerciali/economici con BAT.")

subsection_title("7.3 Modulo Gestione Vendite — Funzionalità Incluse MVP")

body_text("F5 — Form unico di vendita: Scansione codice myglow tramite fotocamera con riconoscimento AI (modello/colore/prezzo). Auto-compilazione operatore e timestamp. Inserimento manuale dati cliente (genere, età). Foto scontrino in-app. Gestione 3 flussi vendita configurabili (scansione obbligatoria, inserimento manuale quantità, vendita Shopify). Gestione prezzi centralizzata con logica sconto clienti registrati BAT. Registrazione in due step per connettività scarsa. Aggiornamento automatico giacenze.")

body_text("F5 — Cross-check automatici: Vendite vs stock (alert vendite mancanti). Inventario periodico anti-furto con confronto codici univoci. Cross-check vendite vs scontrini (se OCR implementato).")

body_text("F6 — Pop-up vendita indiretta: Tipologia pop-up configurabile (vendita diretta vs tabaccaio). Tracciamento legame vendita NOT e tabaccaio. Cambio tipologia pop-up nel tempo.")

body_text("T — Timbratura digitale: Check-in/check-out con geolocalizzazione. Conteggio ore automatico per operatore/giorno/pop-up. Correlazione vendite/ora.")

subsection_title("7.4 Modulo Vendite — Escluso / Future Option")

body_text("Calendario turni collettivo con prenotazione slot (quotato separatamente come modulo opzionale). Export monte ore per studio paghe in formato compatibile (nice-to-have, valutabile in Fase 2). Integrazione API ufficiale con BAT/Pard/Shopify (vedi sezione integrazioni). Modulo formazione operatori (in roadmap, non in scope MVP).")

subsection_title("7.5 Dashboard & Analytics (trasversale)")

body_text("Stock real-time per pop-up, magazzino, prodotto/colore con selettore per tipo pop-up. Previsione rottura stock (indicativa, basata su modello predittivo). Segnalazione anomalie consumo gadget vs vendite. Performance per pop-up e venditore (vendite/ora, trend, confronto). Visualizzazione e regolazione parametri modello predittivo. Export Excel/CSV con filtri. Report settimanale BAT auto-generato.")

subsection_title("7.6 Deliverable Tangibili")

make_ec_table(
    ["Deliverable", "Fase", "Formato"],
    [
        ["Backlog prioritizzato e requisiti di dettaglio", "Discovery", "Documento + tool di PM"],
        ["Architettura tecnica e design di alto livello", "Discovery", "Documento tecnico"],
        ["Wireframe/mockup interfacce principali", "Discovery/Dev", "Figma o equivalente"],
        ["Piattaforma funzionante (moduli Logistica + Vendite)", "Go-Live", "Web App (PWA)"],
        ["Modello AI predittivo addestrato su storico 5 anni", "Go-Live", "Componente integrato"],
        ["Documentazione utente e guida operativa", "Go-Live", "PDF / in-app"],
        ["Training iniziale per Admin e key user (2 sessioni)", "Go-Live", "Sessione live"],
        ["Report di collaudo e test eseguiti", "Go-Live", "Documento"],
    ],
    [8.5, 3.5, 5.5]
)

# ════════════════════════════════════════════════════════════════
# 8. DETTAGLIO FASI
# ════════════════════════════════════════════════════════════════
section_title("8. Dettaglio per Fasi (MVP → Go-Live → Hypercare)")

subsection_title("8.1 Fase 1A — Discovery & Design (Settimane 1-4)")

bold_value("Obiettivo: ", "consolidare i requisiti, definire l'architettura e costruire il backlog di sviluppo.")
bold_value("Contenuti: ", "Process mapping dettagliato dei flussi logistica e vendite. Definizione struttura dati (catalogo prodotti, tipologie pop-up, ruoli). Design architettura multi-tenant e integrazioni. Wireframe interfacce principali (form vendita, dashboard, vista magazziniere). Backlog refinement e sprint planning.")
bold_value("Exit criteria: ", "Backlog approvato da NOT. Architettura validata dal team tecnico. Wireframe confermati da Carmen/Laila.")

subsection_title("8.2 Fase 1B — Sviluppo Core (Settimane 5-14)")

bold_value("Obiettivo: ", "sviluppo e test dei moduli core (Logistica + Vendite + Dashboard).")
bold_value("Contenuti: ", "Sprint 1-2 (sett. 5-8): Setup infrastruttura, autenticazione, gestione prodotti e pop-up, setup iniziale (F0). Sprint 3-4 (sett. 9-12): Form vendita con scansione AI, registrazione vendite, check-in/check-out magazzino e pop-up (F3-F5). Sprint 5 (sett. 13-14): Modello predittivo, workflow approvazione, cross-check automatici, timbratura digitale (F1-F2).")
bold_value("Exit criteria: ", "Funzionalità core testabili in ambiente di staging. Modello predittivo addestrato su storico. Demo intermedia a stakeholder NOT.")

subsection_title("8.3 Fase 1C — Dashboard, Test & Go-Live (Settimane 15-20)")

bold_value("Obiettivo: ", "completamento dashboard, UAT, fine-tuning e rilascio in produzione.")
bold_value("Contenuti: ", "Sprint 6 (sett. 15-16): Dashboard analytics, report, alert predittivi. Sprint 7 (sett. 17-18): UAT con Carmen/Laila e Area Manager pilota. Bug fixing e fine-tuning modello. Sprint 8 (sett. 19-20): Deployment produzione, training utenti, go-live su 3-5 pop-up pilota.")
bold_value("Exit criteria: ", "UAT superata con sign-off Carmen. Go-live su pop-up pilota senza bloccanti. Training completato per Admin e key user.")

subsection_title("8.4 Hypercare & Stabilizzazione (Settimane 21-24)")

bold_value("Obiettivo: ", "supporto post-go-live, rollout progressivo e stabilizzazione.")
bold_value("Contenuti: ", "Supporto operativo dedicato (8 ore/settimana). Monitoraggio performance e fix prioritari. Rollout progressivo sui restanti pop-up. Fine-tuning modello predittivo su dati reali. Raccolta feedback per pianificazione Fase 2.")
bold_value("Exit criteria: ", "Sistema stabile su tutti i pop-up attivi. Nessun bug bloccante aperto da oltre 48 ore. Handover operativo completato.")

# ════════════════════════════════════════════════════════════════
# 9. WBS E STIMA EFFORT
# ════════════════════════════════════════════════════════════════
section_title("9. Team e Stima Effort")

body_text("Il team di sviluppo è composto da due Senior Developer con oltre 10 anni di esperienza ciascuno nella progettazione e realizzazione di soluzioni software custom:")

bold_value("Technical Lead & System Architect — ", "Esperienza consolidata come CTO in una tech company italiana, specializzato in architetture cloud-native, system integration, piattaforme SaaS multi-tenant e real-time data processing. Competenze specifiche in API design, microservizi e infrastrutture scalabili.")

bold_value("AI & Data Engineer — ", "Consulente specializzato in Intelligenza Artificiale e Data Science con focus su machine learning, algoritmi predittivi e sistemi di recommendation. Competenze in Python, modelli di computer vision, NLP e integrazione LLM.")

bold_value("Project Manager (HeyAI) — ", "Coordinamento progettuale, raccolta requisiti, testing, documentazione, training e change management.")

empty_para()

subsection_title("9.1 WBS — Pacchetto Recommended")

make_ec_table(
    ["Workpackage", "Ruoli", "GG/Uomo", "Note"],
    [
        [{"text": "DISCOVERY & DESIGN", "bold": True, "bg": "D9E2F3"}, {"text": "", "bg": "D9E2F3"}, {"text": "25", "bold": True, "bg": "D9E2F3"}, {"text": "", "bg": "D9E2F3"}],
        ["Process mapping, analisi requisiti dettaglio", "PM, BA", "8", "Include interviste stakeholder"],
        ["Architettura tecnica e design sistema", "Tech Lead, AI Eng", "8", "Multi-tenant, data model"],
        ["UX/UI wireframe e prototipazione", "PM, Dev", "5", "Form vendita, dashboard, viste ruolo"],
        ["Backlog refinement e sprint planning", "PM, Tech Lead", "4", "Prioritizzazione con NOT"],
        [{"text": "SVILUPPO CORE — LOGISTICA", "bold": True, "bg": "D9E2F3"}, {"text": "", "bg": "D9E2F3"}, {"text": "45", "bold": True, "bg": "D9E2F3"}, {"text": "", "bg": "D9E2F3"}],
        ["Setup iniziale (F0): censimento, import storico", "Dev x2", "6", "Qualità dati a carico cliente"],
        ["Modello predittivo AI (F1): training, output, parametri", "AI Eng, Dev", "14", "Dipende da qualità storico"],
        ["Workflow approvazione e task magazziniere (F2)", "Dev x2", "8", ""],
        ["Check-out magazzino + check-in merce BAT (F3)", "Dev x2", "8", "Conferma aggregate"],
        ["Check-in pop-up + registrazione codici (F4)", "Dev x2", "6", "Myglow manuale + alert"],
        ["Gestione multi-tenant e categorie prodotto", "Tech Lead, Dev", "3", "Layer filtraggio per tipo pop-up"],
        [{"text": "SVILUPPO CORE — VENDITE", "bold": True, "bg": "D9E2F3"}, {"text": "", "bg": "D9E2F3"}, {"text": "55", "bold": True, "bg": "D9E2F3"}, {"text": "", "bg": "D9E2F3"}],
        ["Form unico vendita con scansione AI (F5.1-F5.9)", "AI Eng, Dev x2", "18", "Computer vision, 3 flussi vendita"],
        ["Gestione catalogo prodotti (back-office)", "Dev", "5", "CRUD, clonazione, categorie"],
        ["Pop-up vendita indiretta (F6)", "Dev", "4", "Configurazione tipo pop-up"],
        ["Cross-check automatici (vendite/stock/incassi)", "Dev, AI Eng", "8", "Alert anomalie"],
        ["Timbratura digitale e correlazione vendite/ora", "Dev", "6", "Geolocalizzazione, calcolo automatico"],
        ["Gestione ruoli, permessi, autenticazione", "Tech Lead", "5", "4 profili, audit base"],
        ["Gestione connettività e re-inserimento", "Dev", "4", "Interfaccia recovery"],
        ["Integrazioni Shopify (K-ippun Haru)", "Dev", "5", "API Shopify confermata"],
        [{"text": "DASHBOARD & ANALYTICS", "bold": True, "bg": "D9E2F3"}, {"text": "", "bg": "D9E2F3"}, {"text": "20", "bold": True, "bg": "D9E2F3"}, {"text": "", "bg": "D9E2F3"}],
        ["Dashboard stock real-time e alert predittivi", "Dev, AI Eng", "8", "Selettore tipo pop-up"],
        ["Dashboard performance venditori e pop-up", "Dev", "6", "KPI, trend, confronto"],
        ["Export report e report BAT automatico", "Dev", "4", "Excel/CSV + template BAT"],
        ["Parametri modello predittivo (visualizzazione)", "Dev", "2", ""],
        [{"text": "TESTING & QUALITY", "bold": True, "bg": "D9E2F3"}, {"text": "", "bg": "D9E2F3"}, {"text": "18", "bold": True, "bg": "D9E2F3"}, {"text": "", "bg": "D9E2F3"}],
        ["Test unitari e di integrazione", "Dev x2", "8", "Copertura funzionalità core"],
        ["UAT con team NOT", "PM, Dev", "6", "2 cicli UAT + fix"],
        ["Security review e performance testing", "Tech Lead", "4", "OWASP base, load test"],
        [{"text": "DEPLOYMENT & GO-LIVE", "bold": True, "bg": "D9E2F3"}, {"text": "", "bg": "D9E2F3"}, {"text": "12", "bold": True, "bg": "D9E2F3"}, {"text": "", "bg": "D9E2F3"}],
        ["Setup infrastruttura e CI/CD", "Tech Lead, DevOps", "4", "Cloud, ambienti staging/prod"],
        ["Deployment e migrazione dati iniziali", "Dev, PM", "3", "Import giacenze e storico"],
        ["Training utenti (2 sessioni)", "PM", "3", "Admin + key user"],
        ["Documentazione utente", "PM", "2", "Guida operativa"],
        [{"text": "PROJECT MANAGEMENT", "bold": True, "bg": "D9E2F3"}, {"text": "", "bg": "D9E2F3"}, {"text": "30", "bold": True, "bg": "D9E2F3"}, {"text": "", "bg": "D9E2F3"}],
        ["Coordinamento team e stakeholder", "PM", "12", "8h/sett x 20 sett = 20gg + buffer"],
        ["Raccolta requisiti, analisi funzionale, UX review", "PM", "8", "Ongoing durante sviluppo"],
        ["QA coordination, documentazione, change mgmt", "PM", "6", ""],
        ["Hypercare post go-live (4 settimane)", "PM, Dev", "4", "8h/sett supporto dedicato"],
        [{"text": "TOTALE PACCHETTO RECOMMENDED", "bold": True, "bg": "1F4E79", "color": WHITE}, {"text": "", "bg": "1F4E79"}, {"text": "205", "bold": True, "bg": "1F4E79", "color": WHITE}, {"text": "", "bg": "1F4E79"}],
    ],
    [6.5, 3.5, 2.0, 5.5]
)

empty_para()

note_text("Il monte ore di sviluppo tecnico viene concordato a forfait con il team di sviluppo e non viene dettagliato per singolo requisito. Lo sviluppo segue un approccio agile con rilasci incrementali.")

# ════════════════════════════════════════════════════════════════
# 10. DIFFERENZE TRA PACCHETTI
# ════════════════════════════════════════════════════════════════
section_title("10. Dettaglio Pacchetti: Base / Recommended / Advanced")

subsection_title("10.1 Pacchetto Base — € 135.000")

body_text("Scope ridotto per validazione del concept con budget contenuto. Include: setup iniziale e import storico, modello predittivo con parametri regolabili (senza ottimizzazione distribuzione avanzata), workflow approvazione e task magazziniere, check-out/check-in base (senza registrazione codici myglow), form vendita semplificato (inserimento manuale, senza scansione AI), dashboard base (stock e vendite, senza analytics predittivi), 4 profili utente, timbratura digitale base (senza geolocalizzazione). Esclusi: scansione AI prodotto, cross-check avanzati, integrazioni esterne, pop-up vendita indiretta, hypercare.")

subsection_title("10.2 Pacchetto Recommended — € 168.000 (consigliato)")

body_text("MVP completo con tutte le funzionalità core confermate nel requirement log. Include tutto il pacchetto Base, più: scansione AI prodotto con computer vision (modello/colore/prezzo), cross-check triplo automatico (vendite/stock/incassi), pop-up vendita indiretta con tracciamento, gestione multi-tenant con categorie prodotto per tipo pop-up, dashboard completa con analytics predittivi e alert rottura stock, integrazione Shopify per K-ippun Haru, timbratura con geolocalizzazione e correlazione vendite/ora, report BAT auto-generato, hypercare 4 settimane. Include training iniziale (2 sessioni) e documentazione operativa.")

subsection_title("10.3 Pacchetto Advanced — € 215.000")

body_text("MVP completo più integrazioni esterne e nice-to-have. Include tutto il pacchetto Recommended, più: integrazione BAT/Pard via API o bot RPA (compilazione automatica form esterni), calendario turni collettivo con prenotazione slot e approvazione AM, OCR scontrino per cross-check automatico vendite/incassi, chat AI per interrogazione dati in linguaggio naturale, export monte ore in formato studio paghe, hypercare estesa (8 settimane).")

# ════════════════════════════════════════════════════════════════
# 11. ASSUNZIONI, ESCLUSIONI, RISCHI
# ════════════════════════════════════════════════════════════════
section_title("11. Assunzioni, Esclusioni, Rischi e Change Control")

subsection_title("11.1 Assunzioni")

bold_value("Dati: ", "NOT fornisce dati storici vendite (5 anni) e giacenze iniziali in formato Excel strutturato entro il kick-off. La qualità e completezza dei dati storici è responsabilità del cliente. I dati sono sufficienti per addestrare il modello predittivo (volume, granularità per prodotto/colore/pop-up).")

bold_value("Stakeholder: ", "Carmen e Laila sono disponibili come product owner per 4-8 ore/settimana durante lo sviluppo. Le approvazioni su wireframe, backlog e UAT avvengono entro 5 giorni lavorativi dalla sottomissione.")

bold_value("Infrastruttura: ", "Il sistema viene deployato su infrastruttura cloud gestita da HeyAI. NOT fornisce accesso ai propri sistemi necessari per le integrazioni incluse nello scope.")

bold_value("Integrazioni: ", "Le API Shopify per K-ippun Haru sono disponibili e documentate. L'integrazione con BAT e Pard NON è inclusa nel pacchetto Recommended (vedi Advanced). Il form vendita nel pacchetto Recommended gestisce i dati per NOT; la compilazione automatica verso piattaforme terze è quotata nel pacchetto Advanced.")

bold_value("Operatività: ", "Ogni pop-up dispone di connettività internet sufficiente (Wi-Fi o rete mobile). I dispositivi degli operatori sono smartphone con fotocamera di qualità adeguata per la scansione prodotti.")

subsection_title("11.2 Esclusioni Esplicite")

make_ec_table(
    ["Esclusione", "Motivazione", "Opzione futura"],
    [
        ["Integrazione API diretta con BAT/Pard", "In attesa verifica disponibilità API da parte di NOT", "Inclusa nel pacchetto Advanced o quotabile separatamente"],
        ["Gestione GDPR / data processing dati clienti finali", "Tema legale in carico a reparto legale Noloop", "Se risolto, il form unificato copre anche dati BAT"],
        ["Scansione barcode a magazzino (prodotto per prodotto)", "Esclusa per decisione operativa presa in call", "Non prevista"],
        ["Integrazione Google Trends", "Evoluzione futura", "Quotabile come add-on"],
        ["Calendario turni collettivo", "Nice-to-have non prioritario", "Incluso nel pacchetto Advanced"],
        ["Formazione operatori strutturata", "NOT gestisce internamente con Train Champion", "Quotabile separatamente"],
        ["Migrazione dati da sistemi diversi da Excel", "Fuori perimetro", "Quotabile su analisi"],
        ["Licenze cloud, device, hardware", "A carico NOT", "Non applicabile"],
        ["Manutenzione integrazioni non ufficiali (RPA/bot)", "Se BAT/Pard non forniscono API", "Fix a pagamento one-shot"],
        ["Modalità offline completa", "Evoluzione futura", "Quotabile come evolutiva Fase 2"],
    ],
    [5.5, 6.0, 6.0]
)

subsection_title("11.3 Rischi Principali e Mitigazioni")

make_ec_table(
    ["Rischio", "Prob.", "Impatto", "Mitigazione"],
    [
        ["Dati storici incompleti o di bassa qualità", "Media", "Alto", "Fase 0 include data audit. Se dati insufficienti, modello parte con regole base e migliora progressivamente"],
        ["Ritardo consegna dati/materiali da NOT", "Alta", "Medio", "Clausola di slittamento timeline proporzionale. Sprint parallelizzabili non bloccati"],
        ["Scope creep (nuovi requisiti durante sviluppo)", "Alta", "Alto", "Change control formale. Backlog congelato per sprint. Nuovi requisiti entrano nel backlog successivo"],
        ["API BAT/Pard non disponibili", "Media", "Medio", "Non impatta pacchetto Recommended. Per Advanced: fallback RPA quotato"],
        ["Resistenza al cambiamento da parte operatori", "Media", "Medio", "Training dedicato. Go-live graduale su pop-up pilota. Feedback loop continuo"],
        ["Performance modello predittivo iniziale", "Bassa", "Medio", "Parametri regolabili. Fine-tuning continuo. Modello migliora con volume dati reali"],
        ["Qualità scansione AI dipendente da device", "Media", "Basso", "Fallback manuale sempre disponibile. Approccio ibrido (CV + OCR + selezione)"],
        ["Complessità multi-tenant superiore alle attese", "Bassa", "Medio", "Architettura validata in Discovery. Buffer 15% incluso nella stima"],
    ],
    [4.5, 1.5, 1.5, 10.0]
)

subsection_title("11.4 Change Control")

body_text("Qualsiasi modifica al perimetro concordato viene gestita tramite il seguente processo:")

body_text("1. NOT formalizza la richiesta di change per iscritto (email o tool di PM).")
body_text("2. HeyAI valuta l'impatto in termini di effort, costi e timeline entro 5 giorni lavorativi.")
body_text("3. Se l'impatto è inferiore al 5% dell'effort totale di fase e non modifica deliverable core, HeyAI assorbe il change nel buffer di progetto.")
body_text("4. Se l'impatto è superiore, HeyAI produce una variante con costi e timeline aggiornati. La variante richiede approvazione scritta prima dell'esecuzione.")
body_text("5. Change request non approvate vengono inserite nel backlog per fasi successive.")

# ════════════════════════════════════════════════════════════════
# 12. GOVERNANCE
# ════════════════════════════════════════════════════════════════
section_title("12. Governance di Progetto")

make_ec_table(
    ["Cerimonia", "Frequenza", "Partecipanti", "Obiettivo"],
    [
        ["Steering Committee", "Mensile", "Germano/Massi (NOT), Carlo (HeyAI)", "Avanzamento, decisioni strategiche, escalation"],
        ["Weekly sync", "Settimanale", "Carmen/Laila (NOT), Carlo (HeyAI)", "Stato sprint, backlog, bloccanti"],
        ["Sprint Review / Demo", "Bi-settimanale", "Team NOT + HeyAI", "Demo funzionalità sviluppate"],
        ["UAT", "Fine Fase 1B e 1C", "Carmen/Laila + AM pilota", "Validazione funzionale"],
        ["Go/No-Go", "Pre go-live", "Germano + Carmen + Carlo", "Decisione rilascio in produzione"],
    ],
    [3.5, 3.0, 5.5, 5.5]
)

# ════════════════════════════════════════════════════════════════
# 13. CONDIZIONI GENERALI
# ════════════════════════════════════════════════════════════════
section_title("13. Condizioni Generali")

empty_para()

make_ec_table(
    ["Condizione", "Dettaglio"],
    [
        [{"text": "Tempi di consegna", "bold": True}, "20 settimane dalla firma del contratto (+ 4 settimane hypercare)"],
        [{"text": "Proprietà intellettuale", "bold": True}, "Codice e algoritmi di proprietà esclusiva del Cliente"],
        [{"text": "Garanzia", "bold": True}, "4 settimane hypercare post-lancio incluse (pacchetto Recommended)"],
        [{"text": "Costi ricorrenti post-MVP", "bold": True}, "Da dettagliare in fase di Discovery (cloud + API AI) a carico Cliente"],
        [{"text": "Change request", "bold": True}, "Modifiche allo scope saranno quotate separatamente (vedi sezione 11.4)"],
        [{"text": "Validità offerta", "bold": True}, "30 giorni dalla data di emissione"],
        [{"text": "Pagamento", "bold": True}, "30 giorni dalla data di fatturazione per ogni milestone"],
        [{"text": "IVA", "bold": True}, "Importi al netto di IVA (22%)"],
    ],
    [5.0, 12.5]
)

# ════════════════════════════════════════════════════════════════
# 14. PUNTI APERTI
# ════════════════════════════════════════════════════════════════
section_title("14. Punti Aperti e Informazioni Mancanti")

subsection_title("14.1 Decisioni Aperte")

make_ec_table(
    ["#", "Punto aperto", "Impatto", "Chi risolve", "Quando"],
    [
        ["1", "API BAT e Pard: disponibilità e documentazione", "Determina se integrazione è API o RPA (delta effort ~15-20 gg)", "NOT verifica con BAT/Pard", "Prima di eventuale upgrade ad Advanced"],
        ["2", "GDPR — transito dati clienti nel form unico", "Se non risolvibile, il form copre solo NOT+Pard (2 piattaforme su 3)", "Legale Noloop", "Prima dello sviluppo del form unificato"],
        ["3", "Lista completa campi dei 3 form attuali", "Dimensiona il form unico e il mapping", "Carmen invia a HeyAI", "Entro kick-off"],
        ["4", "Gestione sostituzione dispositivi difettosi (F5.10)", "Se richiede logiche complesse, potrebbe generare effort aggiuntivo", "Da definire in call tecnica", "Entro fine Discovery"],
        ["5", "Tracciamento NEO/VELO senza codice univoco (F5.12)", "Logica specifica da definire per prodotti monopolio di Stato", "Da definire in call tecnica", "Entro fine Discovery"],
        ["6", "Foto packaging e posizione codici per training modello CV", "Serve per calibrare effort riconoscimento prodotto", "NOT invia foto", "Entro settimana 2"],
        ["7", "Connettività nei centri commerciali: Wi-Fi disponibile?", "Impatta architettura offline/recovery", "NOT verifica", "Entro kick-off"],
        ["8", "Catalogo prodotti completo con casistiche", "Dimensiona flussi vendita configurabili", "NOT invia", "Entro kick-off"],
    ],
    [1.0, 5.0, 4.5, 3.5, 3.5]
)

subsection_title("14.2 Assunzioni che Impattano il Prezzo")

body_text("Se una delle seguenti assunzioni si rivela non valida, il prezzo potrebbe variare:")

make_ec_table(
    ["Assunzione", "Se non valida", "Impatto stimato"],
    [
        ["Dati storici in formato Excel strutturato e pulito", "Serve fase di data cleaning/normalizzazione", "+€ 5.000-10.000"],
        ["4 profili utente standard sufficienti", "Profili aggiuntivi o permessi granulari custom", "+€ 3.000-5.000"],
        ["Catalogo prodotti con massimo 50-80 SKU attivi", "SKU > 100 con logiche complesse", "+€ 2.000-4.000"],
        ["Massimo 2-3 tipologie di pop-up", "Più tipologie con regole dedicate", "+€ 3.000-6.000"],
        ["API Shopify documentata e stabile", "Integrazione Shopify non standard", "+€ 3.000-5.000"],
        ["Scansione AI con 7-8 foto per prodotto sufficienti", "Modello richiede training estensivo", "+€ 5.000-8.000"],
    ],
    [6.0, 5.5, 6.0]
)

# ── Save ──
output_path = "/home/claude/Quotazione_NOT_v2.docx"
doc.save(output_path)
print(f"Documento salvato: {output_path}")